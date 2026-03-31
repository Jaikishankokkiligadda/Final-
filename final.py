import os
import io
import re
import time
import datetime
import streamlit as st


# ── Optional imports ──
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False

try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    try:
        from langchain_community.embeddings import HuggingFaceEmbeddings
    except ImportError:
        HuggingFaceEmbeddings = None

try:
    from langchain_community.vectorstores import FAISS
except ImportError:
    try:
        from langchain.vectorstores import FAISS
    except ImportError:
        FAISS = None

try:
    from langchain_groq import ChatGroq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False
    ChatGroq = None

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ══════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════
st.set_page_config(
    page_title="Indian Lawyer – Satyameva Jayate",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════
# CUSTOM CSS
# ══════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=EB+Garamond:ital,wght@0,400;0,500;1,400&family=DM+Mono:wght@300;400&display=swap');

:root {
  --gold:          #C9A84C;
  --gold-bright:   #E2C06A;
  --gold-deep:     #9B7730;
  --gold-glow:     rgba(201,168,76,0.08);
  --gold-rule:     rgba(201,168,76,0.20);
  --gold-rule-hi:  rgba(201,168,76,0.42);
  --navy:          #080E1C;
  --navy-mid:      #0D1526;
  --navy-surface:  #111C33;
  --navy-raised:   #162040;
  --navy-border:   rgba(255,255,255,0.06);
  --navy-border-hi:rgba(255,255,255,0.12);
  --text-bright:   #F2EEE6;
  --text-mid:      #A09880;
  --text-faint:    #525060;
  --serif: 'EB Garamond', Georgia, serif;
  --mono:  'DM Mono', monospace;
}

html, body, [class*="css"], .stApp {
  font-family: var(--serif) !important;
  background-color: var(--navy) !important;
  color: var(--text-bright) !important;
  -webkit-font-smoothing: antialiased;
}
.block-container { padding: 2rem 2.4rem 3rem !important; max-width: 100% !important; }

section[data-testid="stSidebar"] {
  background: var(--navy-mid) !important;
  border-right: 1px solid var(--gold-rule-hi) !important;
}
section[data-testid="stSidebar"] > div { padding: 1.8rem 1.4rem !important; }

.wm-banner {
  display: flex;
  align-items: center;
  gap: 1.2rem;
  padding: 1.5rem 2rem 1.1rem;
  background: var(--navy-mid);
  border-bottom: 1px solid var(--gold-rule-hi);
  margin: -2rem -2.4rem 2.2rem;
}
.wm-emblem {
  width: 34px; height: 34px;
  border: 1px solid var(--gold-rule-hi);
  border-radius: 50%;
  display: flex; align-items: center; justify-content: center;
  font-size: 1rem; color: var(--gold); flex-shrink: 0;
}
.wm-title { font-family: var(--serif); font-size: 1.3rem; color: var(--text-bright); letter-spacing: 0.04em; }
.wm-title em { font-style: italic; color: var(--gold); }
.wm-rule { flex: 1; height: 1px; background: var(--gold-rule); }
.wm-sanskrit { font-family: var(--serif); font-size: 0.82rem; color: var(--text-faint); letter-spacing: 0.08em; }

.section-eyebrow {
  font-family: var(--mono) !important;
  font-size: 0.6rem !important;
  font-weight: 300 !important;
  letter-spacing: 0.22em !important;
  text-transform: uppercase !important;
  color: var(--gold-deep) !important;
  margin-bottom: 0.3rem !important;
}
.section-title {
  font-family: var(--serif) !important;
  font-size: 1.45rem !important;
  font-weight: 400 !important;
  color: var(--text-bright) !important;
  border-bottom: 1px solid var(--gold-rule) !important;
  padding-bottom: 0.9rem !important;
  margin-bottom: 1.6rem !important;
}

.nav-label {
  font-family: var(--mono);
  font-size: 0.59rem;
  font-weight: 300;
  letter-spacing: 0.22em;
  text-transform: uppercase;
  color: var(--text-faint);
  margin-bottom: 0.5rem;
  display: block;
}
.sidebar-wordmark { padding-bottom: 1.2rem; border-bottom: 1px solid var(--gold-rule); margin-bottom: 1.4rem; }
.sb-title { font-family: var(--serif); font-size: 1.15rem; color: var(--text-bright); letter-spacing: 0.04em; }
.sb-title em { font-style: italic; color: var(--gold); }
.sb-sub { font-family: var(--mono); font-size: 0.62rem; font-weight: 300; color: var(--text-faint); margin-top: 2px; }

.msg-user {
  padding: 0.85rem 1rem;
  border-left: 1.5px solid var(--navy-border-hi);
  margin: 1.1rem 0 1.1rem 12%;
  color: var(--text-mid);
  line-height: 1.78;
  font-size: 1rem;
}
.msg-ai {
  background: var(--gold-glow);
  border-left: 2px solid var(--gold);
  border-radius: 0 4px 4px 0;
  padding: 0.9rem 1.1rem;
  margin: 1.1rem 12% 1.1rem 0;
  color: var(--text-bright);
  line-height: 1.8;
  font-size: 1rem;
}
.msg-role-user { font-family: var(--mono); font-size: 0.6rem; font-weight: 300; letter-spacing: 0.15em; text-transform: uppercase; color: var(--text-faint); margin-bottom: 5px; }
.msg-role-ai   { font-family: var(--mono); font-size: 0.6rem; font-weight: 300; letter-spacing: 0.15em; text-transform: uppercase; color: var(--gold-deep); margin-bottom: 5px; }
.msg-meta { font-family: var(--mono); font-size: 0.6rem; font-weight: 300; color: var(--text-faint); margin-top: 8px; display: flex; gap: 12px; flex-wrap: wrap; }
.src-tag { font-family: var(--mono); font-size: 0.63rem; font-weight: 300; color: #5AB87A; border: 1px solid rgba(91,184,122,0.25); padding: 1px 8px; border-radius: 20px; margin-right: 4px; }

.chip { display:inline-flex; align-items:center; gap:5px; padding:3px 10px; border-radius:20px; font-family:var(--mono); font-size:0.68rem; font-weight:300; }
.chip-ok   { background:rgba(90,184,122,0.08); border:1px solid rgba(90,184,122,0.28); color:#5AB87A; }
.chip-err  { background:rgba(191,107,107,0.08); border:1px solid rgba(191,107,107,0.28); color:#BF6B6B; }
.chip-warn { background:rgba(201,168,76,0.08); border:1px solid var(--gold-rule); color:var(--gold); }

.doc-preview {
  border: 1px solid var(--gold-rule-hi);
  border-radius: 5px;
  padding: 1.8rem 2rem;
  background: var(--navy-surface);
  font-family: var(--serif);
  font-size: 0.88rem;
  line-height: 1.9;
  color: var(--text-bright);
  white-space: pre-wrap;
  max-height: 440px;
  overflow-y: auto;
}

.ref-wrap { margin-bottom: 2rem; }
.ref-head { font-family: var(--mono); font-size: 0.6rem; font-weight: 300; color: var(--gold); letter-spacing: 0.18em; text-transform: uppercase; border-bottom: 1px solid var(--gold-rule-hi); padding-bottom: 0.4rem; margin-bottom: 0.75rem; }
.ref-table { width: 100%; border-collapse: collapse; }
.ref-table tr { border-bottom: 1px solid var(--navy-border); }
.ref-table tr:last-child { border-bottom: none; }
.ref-table td { padding: 0.58rem 0.4rem; font-size: 0.9rem; vertical-align: top; }
.ref-table td:first-child { font-family: var(--mono); font-size: 0.7rem; font-weight: 300; color: var(--gold); width: 210px; padding-right: 1.4rem; white-space: nowrap; }
.ref-table td:last-child { color: var(--text-mid); }

.metric-card { background: var(--navy-surface); border: 1px solid var(--gold-rule); border-radius: 5px; padding: 1rem 1.1rem; text-align: center; }
.metric-val { font-family: var(--serif); font-size: 1.8rem; color: var(--gold); line-height: 1; margin-bottom: 4px; }
.metric-lbl { font-family: var(--mono); font-size: 0.58rem; font-weight: 300; letter-spacing: 0.14em; text-transform: uppercase; color: var(--text-faint); }
.log-row { display: flex; justify-content: space-between; align-items: center; padding: 5px 0; border-bottom: 1px solid var(--navy-border); font-family: var(--mono); font-size: 0.68rem; font-weight: 300; }
.log-row:last-child { border-bottom: none; }
.perf-bar { background: var(--navy-raised); border-radius: 3px; height: 3px; margin-top: 4px; overflow: hidden; }
.perf-fill { height: 100%; border-radius: 3px; background: linear-gradient(90deg, var(--gold-deep), #5AB87A); }

.lawyer-row { display: grid; grid-template-columns: 44px 1fr auto; gap: 0.85rem 1rem; padding: 1.2rem 0.5rem; border-bottom: 1px solid var(--navy-border); transition: background 0.12s, padding 0.12s; border-radius: 4px; }
.lawyer-row:hover { background: var(--gold-glow); padding-left: 0.8rem; }
.l-avatar { width: 42px; height: 42px; border-radius: 50%; background: var(--navy-raised); border: 1px solid var(--gold-rule-hi); display: flex; align-items: center; justify-content: center; font-family: var(--serif); font-size: 0.78rem; color: var(--gold); }
.l-name  { font-size: 0.95rem; font-weight: 500; color: var(--text-bright); display: block; }
.l-role  { font-size: 0.8rem; color: var(--text-mid); font-style: italic; display: block; margin-bottom: 4px; }
.l-tags  { display: flex; gap: 5px; flex-wrap: wrap; }
.l-tag   { font-family: var(--mono); font-size: 0.57rem; font-weight: 300; color: var(--gold-deep); border: 1px solid var(--gold-rule); padding: 1px 7px; border-radius: 20px; }
.l-stat  { text-align: right; }
.l-val   { font-family: var(--serif); font-size: 0.98rem; color: var(--gold); display: block; }
.l-sub   { font-family: var(--mono); font-size: 0.58rem; font-weight: 300; color: var(--text-faint); display: block; }

.doc-grid { display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 1px; background: var(--gold-rule); border: 1px solid var(--gold-rule-hi); border-radius: 5px; overflow: hidden; margin-bottom: 1.8rem; }
.doc-tile { background: var(--navy-surface); padding: 1.2rem 1.1rem; text-align: left; cursor: pointer; transition: background 0.12s; border: none; }
.doc-tile:hover { background: var(--navy-raised); }
.doc-tile-name { font-family: var(--serif); font-size: 0.93rem; color: var(--text-bright); display: block; margin-bottom: 3px; }
.doc-tile-desc { font-size: 0.74rem; color: var(--text-faint); font-style: italic; line-height: 1.4; }

.divider { border: none; border-top: 1px solid var(--gold-rule); margin: 1.2rem 0; }

.stTextInput input, .stTextArea textarea, .stSelectbox select {
  background: var(--navy-surface) !important;
  color: var(--text-bright) !important;
  border: none !important;
  border-bottom: 1px solid var(--navy-border-hi) !important;
  border-radius: 0 !important;
  font-family: var(--serif) !important;
  font-size: 0.95rem !important;
  outline: none !important;
  box-shadow: none !important;
}
.stTextInput input:focus, .stTextArea textarea:focus { border-bottom-color: var(--gold) !important; box-shadow: none !important; }
.stTextInput input::placeholder, .stTextArea textarea::placeholder { color: var(--text-faint) !important; font-style: italic !important; }
label[data-testid="stWidgetLabel"] p { font-family: var(--mono) !important; font-size: 0.6rem !important; font-weight: 300 !important; letter-spacing: 0.14em !important; text-transform: uppercase !important; color: var(--gold-deep) !important; }

.stButton > button { font-family: var(--mono) !important; font-size: 0.68rem !important; font-weight: 400 !important; letter-spacing: 0.1em !important; text-transform: uppercase !important; color: var(--navy) !important; background: var(--gold) !important; border: none !important; border-radius: 2px !important; padding: 0.5rem 1.4rem !important; transition: background 0.15s !important; }
.stButton > button:hover { background: var(--gold-bright) !important; }
.stDownloadButton > button { background: transparent !important; color: var(--text-mid) !important; border: 1px solid var(--navy-border-hi) !important; font-family: var(--mono) !important; font-size: 0.66rem !important; font-weight: 300 !important; }
.stDownloadButton > button:hover { color: var(--gold) !important; border-color: var(--gold-rule-hi) !important; }

.stTabs [data-baseweb="tab-list"] { background: var(--navy-mid) !important; border-bottom: 1px solid var(--gold-rule) !important; gap: 0 !important; }
.stTabs [data-baseweb="tab"] { font-family: var(--mono) !important; font-size: 0.67rem !important; font-weight: 300 !important; letter-spacing: 0.1em !important; text-transform: uppercase !important; color: var(--text-faint) !important; padding: 0.7rem 1.3rem !important; background: transparent !important; }
.stTabs [aria-selected="true"] { color: var(--gold) !important; border-bottom: 2px solid var(--gold) !important; background: transparent !important; }
.stTabs [data-baseweb="tab-highlight"] { display: none !important; }

div[data-baseweb="select"] > div { background: var(--navy-surface) !important; border: 1px solid var(--navy-border-hi) !important; border-radius: 3px !important; font-family: var(--mono) !important; font-size: 0.67rem !important; font-weight: 300 !important; color: var(--text-mid) !important; }
div[data-baseweb="select"] > div:focus-within { border-color: var(--gold-rule-hi) !important; }

details { background: var(--navy-surface) !important; border: 1px solid var(--navy-border) !important; border-radius: 4px !important; margin-bottom: 8px !important; }
details summary { font-family: var(--mono) !important; font-size: 0.67rem !important; font-weight: 300 !important; letter-spacing: 0.12em !important; text-transform: uppercase !important; color: var(--gold) !important; padding: 0.7rem 1rem !important; }
.streamlit-expanderHeader { color: var(--gold) !important; font-family: var(--mono) !important; font-size: 0.67rem !important; }
.streamlit-expanderHeader:hover { color: var(--gold-bright) !important; }

.stSlider [data-baseweb="slider"] { padding: 0 !important; }
.stSlider [data-baseweb="thumb"] { background: var(--gold) !important; }
.stSlider [data-baseweb="track-fill"] { background: var(--gold) !important; }

.stRadio label { font-family: var(--mono) !important; font-size: 0.67rem !important; font-weight: 300 !important; color: var(--text-mid) !important; }

.stAlert { background: var(--navy-surface) !important; border: 1px solid var(--navy-border-hi) !important; border-radius: 4px !important; }
div[data-testid="stNotification"] { background: var(--navy-surface) !important; }

::-webkit-scrollbar { width: 3px; height: 3px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: var(--navy-raised); border-radius: 2px; }
::-webkit-scrollbar-thumb:hover { background: var(--gold-deep); }

.empty-state { text-align: center; padding: 3.5rem 0 2.5rem; }
.empty-glyph { font-size: 2.2rem; color: var(--navy-raised); }
.empty-title { font-family: var(--serif); font-size: 1.25rem; color: var(--text-bright); margin: 0.6rem 0 0.4rem; }
.empty-sub { font-size: 0.9rem; color: var(--text-mid); max-width: 400px; margin: 0 auto; line-height: 1.6; }

.info-note { font-family: var(--mono); font-size: 0.7rem; font-weight: 300; color: var(--text-faint); font-style: italic; text-align: center; padding: 0.8rem 0; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════
# CONSTANTS
# ══════════════════════════════════════════════
VECTORSTORE_PATH = "vectorstore"
EMBED_MODEL      = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_SIZE       = 800
CHUNK_OVERLAP    = 100

GROQ_MODELS = {
    "Llama 3.1 8B — Fast":    "llama-3.1-8b-instant",
    "Llama 3.3 70B — Smart":  "llama-3.3-70b-versatile",
    "Mixtral 8×7B — Balanced":"mixtral-8x7b-32768",
    "Gemma 2 9B":              "gemma2-9b-it",
}

DOCUMENT_TEMPLATES = {
    "FIR Draft": {
        "icon": "📋", "desc": "First Information Report to file with police",
        "fields": ["complainant_name","complainant_address","complainant_phone",
                   "incident_date","incident_place","accused_name",
                   "incident_description","witnesses","police_station"],
        "labels": {
            "complainant_name":    "Complainant Full Name",
            "complainant_address": "Complainant Address",
            "complainant_phone":   "Complainant Phone",
            "incident_date":       "Date of Incident",
            "incident_place":      "Place of Incident",
            "accused_name":        "Accused Name(s)",
            "incident_description":"Describe the Incident in Detail",
            "witnesses":           "Witness Names (if any)",
            "police_station":      "Police Station Name & District",
        },
        "textarea_fields": ["complainant_address","incident_description"],
        "prompt": lambda fields, today: f"""Draft a formal FIR for Indian police, dated {today}.
Include: station header, FIR number placeholder, complainant details, detailed incident narration, applicable IPC/BNS sections, prayer for action.
Details:
{fields}
Write the complete FIR:""",
    },
    "Legal Notice": {
        "icon": "📜", "desc": "Formal legal notice to another party",
        "fields": ["sender_name","sender_address","sender_advocate","recipient_name",
                   "recipient_address","notice_subject","facts","demand","reply_days"],
        "labels": {
            "sender_name":      "Sender / Client Full Name",
            "sender_address":   "Sender Address",
            "sender_advocate":  "Advocate Name (optional)",
            "recipient_name":   "Recipient / Opposite Party Name",
            "recipient_address":"Recipient Address",
            "notice_subject":   "Subject of Notice",
            "facts":            "Facts & Circumstances",
            "demand":           "Demand / Relief Sought",
            "reply_days":       "Days to Reply (e.g. 15, 30)",
        },
        "textarea_fields": ["sender_address","recipient_address","facts","demand"],
        "prompt": lambda fields, today: f"""Draft a formal Legal Notice under Indian law, dated {today}.
Include: LEGAL NOTICE heading, sender/advocate details, recipient details, numbered facts, legal grounds, specific demand with timeline, consequences of non-compliance.
Details:
{fields}
Write the complete Legal Notice:""",
    },
    "Bail Application": {
        "icon": "🔓", "desc": "Application for regular or anticipatory bail",
        "fields": ["applicant_name","applicant_age","applicant_address","fir_number",
                   "police_station","sections_charged","arrest_date","court_name","bail_grounds"],
        "labels": {
            "applicant_name":    "Applicant / Accused Full Name",
            "applicant_age":     "Age of Applicant",
            "applicant_address": "Applicant's Address",
            "fir_number":        "FIR Number",
            "police_station":    "Police Station",
            "sections_charged":  "IPC/Other Sections Charged",
            "arrest_date":       "Date of Arrest",
            "court_name":        "Court Name",
            "bail_grounds":      "Grounds for Bail",
        },
        "textarea_fields": ["applicant_address","bail_grounds"],
        "prompt": lambda fields, today: f"""Draft a Bail Application under Section 437/439 CrPC (or BNSS equivalent) for an Indian court, dated {today}.
Include: court heading, applicant details, FIR details, at least 5 specific grounds for bail, prayer clause.
Details:
{fields}
Write the complete Bail Application:""",
    },
    "Affidavit": {
        "icon": "📝", "desc": "General purpose sworn affidavit",
        "fields": ["deponent_name","deponent_age","deponent_address","deponent_occupation",
                   "affidavit_subject","statement_content","place","affidavit_date"],
        "labels": {
            "deponent_name":      "Deponent Full Name",
            "deponent_age":       "Age",
            "deponent_address":   "Address",
            "deponent_occupation":"Occupation",
            "affidavit_subject":  "Subject / Purpose of Affidavit",
            "statement_content":  "Content / Statements to Declare",
            "place":              "Place of Execution",
            "affidavit_date":     "Date",
        },
        "textarea_fields": ["deponent_address","statement_content"],
        "prompt": lambda fields, today: f"""Draft a formal Affidavit under Indian law, dated {today}.
Include: heading, court/authority details, deponent details, numbered factual statements, solemn declaration, verification clause, signature block.
Details:
{fields}
Write the complete Affidavit:""",
    },
    "Rent Agreement": {
        "icon": "🏠", "desc": "Residential / commercial rent agreement draft",
        "fields": ["landlord_name","landlord_address","tenant_name","tenant_address",
                   "property_address","rent_amount","security_deposit",
                   "lease_start","lease_duration","special_terms"],
        "labels": {
            "landlord_name":     "Landlord Full Name",
            "landlord_address":  "Landlord Address",
            "tenant_name":       "Tenant Full Name",
            "tenant_address":    "Tenant Address",
            "property_address":  "Property / Premises Address",
            "rent_amount":       "Monthly Rent (₹)",
            "security_deposit":  "Security Deposit (₹)",
            "lease_start":       "Lease Start Date",
            "lease_duration":    "Lease Duration (e.g. 11 months)",
            "special_terms":     "Special Terms & Conditions (optional)",
        },
        "textarea_fields": ["landlord_address","tenant_address","property_address","special_terms"],
        "prompt": lambda fields, today: f"""Draft a comprehensive Rental Agreement under Indian law, dated {today}.
Include: parties, property description, at least 10 detailed clauses covering rent, maintenance, termination, lock-in period, dispute resolution, signature block with witnesses.
Details:
{fields}
Write the complete Rent Agreement:""",
    },
    "Consumer Complaint": {
        "icon": "🛒", "desc": "Consumer forum complaint against seller/service",
        "fields": ["complainant_name","complainant_address","complainant_phone",
                   "opposite_party_name","opposite_party_address","purchase_date",
                   "product_service","complaint_details","relief_sought","forum_name"],
        "labels": {
            "complainant_name":      "Complainant Full Name",
            "complainant_address":   "Complainant Address",
            "complainant_phone":     "Phone Number",
            "opposite_party_name":   "Opposite Party / Company Name",
            "opposite_party_address":"Opposite Party Address",
            "purchase_date":         "Date of Purchase / Service",
            "product_service":       "Product / Service Name",
            "complaint_details":     "Details of Complaint / Deficiency",
            "relief_sought":         "Relief / Compensation Sought",
            "forum_name":            "Consumer Forum Name & District",
        },
        "textarea_fields": ["complainant_address","opposite_party_address","complaint_details","relief_sought"],
        "prompt": lambda fields, today: f"""Draft a Consumer Complaint under the Consumer Protection Act 2019, dated {today}.
Include: forum heading, complainant and opposite party details, jurisdiction, numbered statement of facts, legal grounds, specific relief sought, verification.
Details:
{fields}
Write the complete Consumer Complaint:""",
    },
}

LAWYERS = [
    {"initials":"PS","name":"Adv. Priya Sharma",  "role":"Senior Advocate — Supreme Court of India","city":"Delhi",    "specs":["Criminal Law","Constitutional Law","Bail & Appeals"],    "fee":"₹5,000","rating":"4.9"},
    {"initials":"RM","name":"Adv. Rajesh Menon",  "role":"High Court Advocate — Bombay High Court", "city":"Mumbai",   "specs":["Corporate Law","Contract Disputes","Cheque Bounce"],    "fee":"₹3,500","rating":"4.7"},
    {"initials":"SR","name":"Adv. Sunita Rao",    "role":"Family Court Specialist — Karnataka HC",  "city":"Bangalore","specs":["Divorce","Child Custody","Domestic Violence"],           "fee":"₹2,500","rating":"4.8"},
    {"initials":"VJ","name":"Adv. Vikram Joshi",  "role":"Revenue & Property — Pune District Court","city":"Pune",     "specs":["Property Law","Rent Disputes","Land Acquisition"],       "fee":"₹2,000","rating":"4.6"},
    {"initials":"FK","name":"Adv. Fatima Khan",   "role":"Labour & Employment — Telangana HC",      "city":"Hyderabad","specs":["Labour Law","Employment Disputes","Wrongful Termination"],"fee":"₹2,000","rating":"4.7"},
    {"initials":"DN","name":"Adv. Deepak Nair",   "role":"Consumer & Civil — Madras High Court",    "city":"Chennai",  "specs":["Consumer Protection","Civil Suits","Medical Negligence"],"fee":"₹1,500","rating":"4.5"},
    {"initials":"AG","name":"Adv. Ananya Gupta",  "role":"Tax & Regulatory — Calcutta High Court",  "city":"Kolkata",  "specs":["Income Tax","GST","Customs & Excise"],                   "fee":"₹4,000","rating":"4.8"},
    {"initials":"AS","name":"Adv. Arjun Singh",   "role":"Criminal Defence — Allahabad High Court", "city":"Lucknow",  "specs":["Criminal Defence","NDPS","POCSO"],                       "fee":"₹2,500","rating":"4.6"},
    {"initials":"MP","name":"Adv. Meera Pillai",  "role":"Cyber Law & IPR — Karnataka HC",          "city":"Bangalore","specs":["Cyber Crime","Intellectual Property","Data Privacy"],    "fee":"₹3,000","rating":"4.9"},
    {"initials":"HB","name":"Adv. Harish Bhatia", "role":"RTI & Administrative Law — Rajasthan HC", "city":"Jaipur",   "specs":["RTI","PILs","Administrative Law"],                       "fee":"₹1,500","rating":"4.4"},
]

# ══════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════
_defaults = {
    "chat_history":       [],
    "vectorstore":        None,
    "active_tab":         "chat",
    "selected_doc":       None,
    "doc_form_data":      {},
    "generated_doc_text": "",
    "generated_doc_bytes":None,
    "metrics": {
        "total_queries":        0,
        "total_doc_generations":0,
        "total_tokens_est":     0,
        "query_log":            [],
        "llm_times":            [],
        "retrieval_times":      [],
        "model_usage":          {},
    },
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ══════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════
def avg(lst):
    return round(sum(lst) / len(lst)) if lst else 0

def estimate_tokens(text):
    return max(1, len(text) // 4)

def get_api_key():
    try:
        return st.secrets["GROQ_API_KEY"].strip()
    except Exception:
        return ""

def log_query(qtype, model, ret_ms, llm_ms, tokens):
    m = st.session_state.metrics
    m["total_queries"]    += 1
    m["total_tokens_est"] += tokens
    m["llm_times"].append(llm_ms)
    m["retrieval_times"].append(ret_ms)
    m["model_usage"][model] = m["model_usage"].get(model, 0) + 1
    m["query_log"].append({
        "time":         datetime.datetime.now().strftime("%H:%M:%S"),
        "type":         qtype,
        "model":        model,
        "retrieval_ms": ret_ms,
        "llm_ms":       llm_ms,
        "total_ms":     ret_ms + llm_ms,
        "tokens":       tokens,
    })
    if len(m["query_log"]) > 50:
        m["query_log"] = m["query_log"][-50:]

# ══════════════════════════════════════════════
# CACHED RESOURCES
# ══════════════════════════════════════════════
@st.cache_resource(show_spinner=False)
def load_embeddings():
    if HuggingFaceEmbeddings is None:
        return None
    return HuggingFaceEmbeddings(
        model_name=EMBED_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"batch_size": 32, "normalize_embeddings": True},
    )

# ── FIX: guard against missing langchain-groq package ──
def load_llm(model_id):
    if not GROQ_AVAILABLE:
        st.error(
            "**langchain-groq is not installed.**  \n"
            "Add `langchain-groq` to your `requirements.txt` and redeploy the app."
        )
        return None
    api_key = get_api_key()
    if not api_key:
        st.error(
            "**GROQ_API_KEY not found.**  \n"
            "Go to your Streamlit Cloud app → **Settings → Secrets** and add:\n"
            "```toml\nGROQ_API_KEY = \"gsk_your_key_here\"\n```"
        )
        return None
    return ChatGroq(
        model=model_id,
        groq_api_key=api_key,
        temperature=0.2,
        max_tokens=1200,
        streaming=True,
    )

def build_vectorstore(files, embeddings):
    if not PDF_AVAILABLE or not LANGCHAIN_AVAILABLE:
        st.error("Install pypdf and langchain: pip install pypdf langchain-text-splitters")
        return None
    docs = []
    for uf in files:
        try:
            reader = PdfReader(uf)
            text = "".join(p.extract_text() or "" for p in reader.pages)
            if text.strip():
                docs.append(Document(page_content=text, metadata={"source": uf.name}))
        except Exception as e:
            st.warning(f"⚠️ {uf.name}: {e}")
    if not docs:
        return None
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks   = splitter.split_documents(docs)
    vs = FAISS.from_documents(chunks, embeddings)
    vs.save_local(VECTORSTORE_PATH)
    return vs

def load_saved_vs(embeddings):
    if FAISS is None or not os.path.exists(VECTORSTORE_PATH):
        return None
    return FAISS.load_local(
        VECTORSTORE_PATH,
        embeddings=embeddings,
        allow_dangerous_deserialization=True,
    )

def chat_prompt(question, context, history):
    hist_txt = ""
    for t in history[-4:]:
        c = t["content"][:500] + "…" if len(t["content"]) > 500 else t["content"]
        hist_txt += f"{'User' if t['role'] == 'user' else 'Assistant'}: {c}\n"
    ctx = context[:2000] + "…" if len(context) > 2000 else context
    return f"""You are Indian Lawyer – Satyameva Jayate, a senior Indian legal expert.

Reply using EXACTLY this structure:

### ⚖️ Legal Explanation
[Clear plain-language explanation]

### 📋 Relevant Sections & Articles
[Key IPC/BNS/Constitutional provisions with section numbers]

### 🏛️ Relevant Case Laws
[1–2 landmark Supreme Court or High Court cases if applicable]

### 💡 Practical Next Steps
[Concrete, actionable steps]

History:
{hist_txt}

Context from uploaded documents:
{ctx}

Question: {question}
Answer:"""

def build_doc_prompt(doc_type, fields):
    today  = datetime.date.today().strftime("%d %B %Y")
    meta   = DOCUMENT_TEMPLATES[doc_type]
    fs_str = "\n".join(f"- {k.replace('_',' ').title()}: {v}" for k, v in fields.items() if v)
    return meta["prompt"](fs_str, today)

def create_docx(title, content):
    doc = DocxDocument()
    s = doc.sections[0]
    s.top_margin    = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin   = Inches(1.25)
    s.right_margin  = Inches(1.25)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(title.upper())
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x5e)
    doc.add_paragraph()

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph(); continue
        para  = doc.add_paragraph()
        clean = line.lstrip("#").strip()
        if line.startswith("###") or (line.isupper() and len(line) > 4):
            r = para.add_run(clean); r.bold = True
            r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x5e)
            para.paragraph_format.space_before = Pt(10)
        elif line.startswith("**") and line.endswith("**"):
            r = para.add_run(line.strip("*")); r.bold = True; r.font.size = Pt(11)
        elif line.startswith("-") or line.startswith("•"):
            para.style = doc.styles["List Bullet"]
            r = para.add_run(line.lstrip("-•").strip()); r.font.size = Pt(11)
        elif re.match(r"^\d+\.", line):
            para.style = doc.styles["List Number"]
            r = para.add_run(re.sub(r"^\d+\.\s*", "", line)); r.font.size = Pt(11)
        else:
            r = para.add_run(clean); r.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ══════════════════════════════════════════════
# WORDMARK BANNER
# ══════════════════════════════════════════════
api_key = get_api_key()
key_chip = (
    '<span class="chip chip-ok">● Groq ready</span>'
    if (api_key and GROQ_AVAILABLE) else
    '<span class="chip chip-err">● langchain-groq not installed — add to requirements.txt</span>'
    if not GROQ_AVAILABLE else
    '<span class="chip chip-err">● GROQ_API_KEY missing — add to Streamlit Secrets</span>'
)

st.markdown(f"""
<div class="wm-banner">
  <div class="wm-emblem">⚖</div>
  <span class="wm-title">Indian Lawyer — <em>Satyameva Jayate</em></span>
  <span class="wm-rule"></span>
  <span class="wm-sanskrit">सत्यमेव जयते</span>
  <span style="margin-left:1rem">{key_chip}</span>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════
with st.sidebar:
    st.markdown("""
    <div class="sidebar-wordmark">
      <div class="sb-title">Indian Lawyer — <em>Satyameva Jayate</em></div>
      <div class="sb-sub">AI Legal Assistant · Groq</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<span class="nav-label">Model</span>', unsafe_allow_html=True)
    model_label = st.selectbox("Model", list(GROQ_MODELS.keys()), label_visibility="collapsed")
    model_id    = GROQ_MODELS[model_label]

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<span class="nav-label">Knowledge Base (RAG)</span>', unsafe_allow_html=True)
    top_k   = st.slider("Chunks to retrieve", 2, 8, 3)
    kb_mode = st.radio("", ["Upload PDFs", "Load saved index"], label_visibility="collapsed")

    if kb_mode == "Upload PDFs":
        uploaded = st.file_uploader("", type="pdf", accept_multiple_files=True, label_visibility="collapsed")
        if st.button("⚡  Build Index", use_container_width=True):
            if not uploaded:
                st.error("Upload at least one PDF.")
            else:
                with st.spinner("Indexing PDFs…"):
                    t0  = time.time()
                    emb = load_embeddings()
                    if emb is None:
                        st.error("langchain-huggingface not installed.")
                    else:
                        vs  = build_vectorstore(uploaded, emb)
                        ms  = round((time.time() - t0) * 1000)
                        if vs:
                            st.session_state.vectorstore = vs
                            st.success(f"Indexed {len(uploaded)} PDF(s) in {ms}ms")
                        else:
                            st.error("No readable text found in uploaded PDFs.")
    else:
        if st.button("📂  Load Saved Index", use_container_width=True):
            with st.spinner("Loading…"):
                emb = load_embeddings()
                vs  = load_saved_vs(emb) if emb else None
                if vs:
                    st.session_state.vectorstore = vs
                    st.success("Loaded saved index.")
                else:
                    st.error("No saved index found. Upload PDFs first.")

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    kb_ok   = st.session_state.vectorstore is not None
    m       = st.session_state.metrics
    avg_ms  = avg(m["llm_times"])
    lat_clr = "#5AB87A" if avg_ms < 3000 else "#C9A84C" if avg_ms < 8000 else "#BF6B6B"
    kb_lbl  = "KB Ready" if kb_ok else "No KB loaded"

    st.markdown(f"""
    <div style="display:flex;flex-direction:column;gap:6px">
      <div style="display:flex;align-items:center;gap:6px;font-family:var(--mono);font-size:0.63rem;font-weight:300;color:var(--text-faint)">
        <span style="width:5px;height:5px;border-radius:50%;flex-shrink:0;background:{'#5AB87A' if kb_ok else 'var(--text-faint)'}"></span>
        {kb_lbl}
      </div>
      <div style="font-family:var(--mono);font-size:0.63rem;font-weight:300;color:var(--text-faint)">
        Model: <span style="color:var(--gold)">{model_label.split('—')[0].strip()}</span>
      </div>
      <div style="font-family:var(--mono);font-size:0.63rem;font-weight:300;color:var(--text-faint)">
        Queries: <span style="color:var(--text-bright)">{m["total_queries"]}</span>
        &nbsp;·&nbsp; ~<span style="color:var(--text-bright)">{m["total_tokens_est"]:,}</span> tokens
      </div>
      {"<div style='font-family:var(--mono);font-size:0.63rem;font-weight:300;color:var(--text-faint)'>Avg LLM: <span style='color:" + lat_clr + "'>" + f"{avg_ms/1000:.1f}s</span></div>" if m["llm_times"] else ""}
    </div>
    """, unsafe_allow_html=True)

# ══════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════
tab_chat, tab_docs, tab_lawyers, tab_ref, tab_metrics = st.tabs([
    "Legal Chat",
    "Document Generator",
    "Find a Lawyer",
    "Reference",
    "Metrics",
])

# ════════════════════════════════════════════════════
# TAB 1 — CHAT
# ════════════════════════════════════════════════════
with tab_chat:
    st.markdown('<p class="section-eyebrow">Legal Chat</p>', unsafe_allow_html=True)
    st.markdown('<h1 class="section-title">Ask any question on Indian law</h1>', unsafe_allow_html=True)

    if not st.session_state.chat_history:
        st.markdown("""
        <div class="empty-state">
          <div class="empty-glyph">⚖</div>
          <p class="empty-title">What is your legal question?</p>
          <p class="empty-sub">Ask about IPC sections, your rights, bail, consumer protection, property, or anything under Indian law.</p>
        </div>
        """, unsafe_allow_html=True)
    else:
        for turn in st.session_state.chat_history:
            if turn["role"] == "user":
                st.markdown(
                    f'<div class="msg-user"><div class="msg-role-user">You</div>{turn["content"]}</div>',
                    unsafe_allow_html=True,
                )
            else:
                ch   = turn["content"].replace("\n", "<br>").replace("### ", "<strong>").replace("**", "")
                meta = turn.get("meta", {})
                meta_html = ""
                if meta:
                    ret = f'<span>🔍 {meta["retrieval_ms"]}ms</span>' if meta.get("retrieval_ms") else ""
                    meta_html = f'<div class="msg-meta"><span>⚡ {meta["llm_ms"]/1000:.1f}s</span>{ret}<span>~{meta["tokens"]} tokens</span></div>'
                srcs = "".join(
                    f'<span class="src-tag">📄 {s}</span>'
                    for s in turn.get("sources", [])
                )
                st.markdown(
                    f'<div class="msg-ai"><div class="msg-role-ai">⚖ Satyameva Jayate</div>'
                    f'{ch}<div style="margin-top:8px">{srcs}</div>{meta_html}</div>',
                    unsafe_allow_html=True,
                )

    question = st.text_area(
        "Your question",
        placeholder="e.g. What are my rights if arrested without a warrant?  ·  What is IPC 420?",
        height=80,
        label_visibility="collapsed",
    )

    c1, c2, c3 = st.columns([3, 1, 1])
    with c1: ask    = st.button("Ask →", use_container_width=True)
    with c2: clear  = st.button("Clear", use_container_width=True)
    with c3: export = st.button("Export", use_container_width=True)

    if clear:
        st.session_state.chat_history = []
        st.rerun()

    if export and st.session_state.chat_history:
        lines = []
        for t in st.session_state.chat_history:
            r = "YOU" if t["role"] == "user" else "SATYAMEVA JAYATE"
            lines.append(f"[{r}]\n{t['content']}\n{'─'*60}")
            if t.get("sources"):
                lines.append(f"Sources: {', '.join(t['sources'])}")
        st.download_button(
            "⬇ Download conversation",
            "\n\n".join(lines),
            f"legal_chat_{datetime.date.today()}.txt",
            "text/plain",
        )

    if ask and question.strip():
        st.session_state.chat_history.append({"role": "user", "content": question.strip()})

        sources      = []
        retrieval_ms = 0
        context      = "No document context loaded. Answer from general Indian legal knowledge."

        if st.session_state.vectorstore:
            with st.spinner("Searching legal documents…"):
                t_ret  = time.time()
                docs   = st.session_state.vectorstore.as_retriever(
                    search_kwargs={"k": top_k}
                ).invoke(question.strip())
                retrieval_ms = round((time.time() - t_ret) * 1000)
            sources = list({d.metadata.get("source", "unknown") for d in docs})
            context = "\n\n".join(
                f"[{d.metadata.get('source','?')}]\n{d.page_content}" for d in docs
            )

        prompt = chat_prompt(question.strip(), context, st.session_state.chat_history[:-1])
        llm    = load_llm(model_id)
        if llm is None:
            st.session_state.chat_history.pop()
            st.stop()

        resp_area = st.empty()
        full_resp = ""
        t_llm     = time.time()

        with st.spinner("Consulting…"):
            for chunk in llm.stream(prompt):
                full_resp += chunk.content
                resp_area.markdown(full_resp + " ▌")

        resp_area.empty()
        llm_ms = round((time.time() - t_llm) * 1000)
        tokens = estimate_tokens(full_resp)
        log_query("chat", model_label, retrieval_ms, llm_ms, tokens)

        st.session_state.chat_history.append({
            "role":    "assistant",
            "content": full_resp,
            "sources": sources,
            "meta":    {"retrieval_ms": retrieval_ms, "llm_ms": llm_ms, "tokens": tokens},
        })
        st.rerun()

# ════════════════════════════════════════════════════
# TAB 2 — DOC GENERATOR
# ════════════════════════════════════════════════════
with tab_docs:
    st.markdown('<p class="section-eyebrow">Document Generator</p>', unsafe_allow_html=True)
    st.markdown('<h1 class="section-title">Draft a legal document</h1>', unsafe_allow_html=True)

    if not st.session_state.selected_doc:
        tiles_html = '<div class="doc-grid">'
        for doc_type, meta in DOCUMENT_TEMPLATES.items():
            tiles_html += (
                f'<div class="doc-tile">'
                f'<span class="doc-tile-name">{meta["icon"]} {doc_type}</span>'
                f'<span class="doc-tile-desc">{meta["desc"]}</span>'
                f'</div>'
            )
        tiles_html += '</div>'
        st.markdown(tiles_html, unsafe_allow_html=True)

        cols = st.columns(3)
        for i, (doc_type, meta) in enumerate(DOCUMENT_TEMPLATES.items()):
            with cols[i % 3]:
                if st.button(f"Select — {doc_type}", key=f"sel_{doc_type}", use_container_width=True):
                    st.session_state.selected_doc        = doc_type
                    st.session_state.doc_form_data       = {}
                    st.session_state.generated_doc_text  = ""
                    st.session_state.generated_doc_bytes = None
                    st.rerun()
    else:
        doc_type = st.session_state.selected_doc
        meta     = DOCUMENT_TEMPLATES[doc_type]

        col_back, col_title = st.columns([1, 7])
        with col_back:
            if st.button("← Back"):
                st.session_state.selected_doc = None
                st.rerun()
        with col_title:
            st.markdown(
                f'<div style="font-family:var(--serif);font-size:1.1rem;color:var(--text-bright);'
                f'padding-top:6px">{meta["icon"]} {doc_type}</div>',
                unsafe_allow_html=True,
            )

        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        left, right = st.columns([1, 1])

        with left:
            st.markdown('<span class="nav-label">Fill in the details</span>', unsafe_allow_html=True)
            form_data = {}
            for field in meta["fields"]:
                label = meta["labels"].get(field, field.replace("_", " ").title())
                key   = f"form_{doc_type}_{field}"
                if field in meta.get("textarea_fields", []):
                    val = st.text_area(label, key=key, height=80)
                else:
                    val = st.text_input(label, key=key)
                form_data[field] = val

            generate_btn = st.button(f"Generate {doc_type} →", use_container_width=True)

        with right:
            st.markdown('<span class="nav-label">Generated Document</span>', unsafe_allow_html=True)
            if generate_btn:
                filled = {k: v for k, v in form_data.items() if v.strip()}
                if len(filled) < 2:
                    st.error("Please fill in at least a few fields first.")
                else:
                    with st.spinner(f"Drafting {doc_type}…"):
                        llm = load_llm(model_id)
                        if llm is None:
                            st.stop()
                        t_doc  = time.time()
                        resp   = llm.invoke(build_doc_prompt(doc_type, filled))
                        text   = resp.content
                        doc_ms = round((time.time() - t_doc) * 1000)
                        tokens = estimate_tokens(text)
                        st.session_state.generated_doc_text = text
                        st.session_state.metrics["total_doc_generations"] += 1
                        log_query("doc_gen", model_label, 0, doc_ms, tokens)
                        if DOCX_AVAILABLE:
                            try:
                                st.session_state.generated_doc_bytes = create_docx(doc_type, text)
                            except Exception as e:
                                st.warning(f"DOCX generation failed: {e}")
                                st.session_state.generated_doc_bytes = None

            if st.session_state.generated_doc_text:
                st.markdown(
                    f'<div class="doc-preview">{st.session_state.generated_doc_text}</div>',
                    unsafe_allow_html=True,
                )
                st.markdown("<br>", unsafe_allow_html=True)
                dl1, dl2 = st.columns(2)
                with dl1:
                    st.download_button(
                        "Download .txt",
                        data=st.session_state.generated_doc_text,
                        file_name=f"{doc_type.lower().replace(' ','_')}.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )
                with dl2:
                    if st.session_state.generated_doc_bytes:
                        st.download_button(
                            "Download .docx",
                            data=st.session_state.generated_doc_bytes,
                            file_name=f"{doc_type.lower().replace(' ','_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                    else:
                        st.markdown(
                            '<div class="info-note">Install python-docx for .docx export</div>',
                            unsafe_allow_html=True,
                        )
            else:
                st.markdown("""
                <div class="empty-state">
                  <div class="empty-glyph">📝</div>
                  <p class="empty-title">Document will appear here</p>
                  <p class="empty-sub">Fill the form and click Generate</p>
                </div>
                """, unsafe_allow_html=True)

# ════════════════════════════════════════════════════
# TAB 3 — FIND A LAWYER
# ════════════════════════════════════════════════════
with tab_lawyers:
    st.markdown('<p class="section-eyebrow">Find a Lawyer</p>', unsafe_allow_html=True)
    st.markdown('<h1 class="section-title">Advocates across India</h1>', unsafe_allow_html=True)

    filter_q = st.text_input(
        "Search",
        placeholder="Filter by city, specialisation, or keyword…",
        label_visibility="collapsed",
    )
    q = filter_q.lower()
    filtered = [
        l for l in LAWYERS
        if not q or q in l["name"].lower() or q in l["city"].lower()
        or any(q in s.lower() for s in l["specs"]) or q in l["role"].lower()
    ] if q else LAWYERS

    rows_html = ""
    for l in filtered:
        tags = "".join(f'<span class="l-tag">{s}</span>' for s in l["specs"])
        tags += f'<span class="l-tag">{l["city"]}</span>'
        rows_html += f"""
        <div class="lawyer-row">
          <div class="l-avatar">{l["initials"]}</div>
          <div>
            <span class="l-name">{l["name"]}</span>
            <span class="l-role">{l["role"]}</span>
            <div class="l-tags">{tags}</div>
          </div>
          <div class="l-stat">
            <span class="l-val">{l["rating"]} ★</span><span class="l-sub">rating</span>
            <span class="l-val" style="font-size:.82rem;margin-top:5px">{l["fee"]}</span><span class="l-sub">consult</span>
          </div>
        </div>"""
    st.markdown(rows_html, unsafe_allow_html=True)

# ════════════════════════════════════════════════════
# TAB 4 — REFERENCE
# ════════════════════════════════════════════════════
with tab_ref:
    st.markdown('<p class="section-eyebrow">Quick Reference</p>', unsafe_allow_html=True)
    st.markdown('<h1 class="section-title">Key Indian Legal Provisions</h1>', unsafe_allow_html=True)

    ref_data = {
        "IPC — Offences Against Person": [
            ("IPC 299 / BNS 100","Culpable Homicide"),
            ("IPC 300 / BNS 101","Murder"),
            ("IPC 302 / BNS 103","Punishment — death or life imprisonment"),
            ("IPC 304B / BNS 80","Dowry Death — minimum 7 years"),
            ("IPC 307 / BNS 109","Attempt to Murder"),
            ("IPC 354 / BNS 74","Assault on woman — 1–5 years"),
            ("IPC 376 / BNS 64","Rape — 7 years to Life"),
        ],
        "IPC — Property & Fraud": [
            ("IPC 378 / BNS 303","Theft"),
            ("IPC 383 / BNS 308","Extortion"),
            ("IPC 406 / BNS 316","Criminal Breach of Trust — 3 years"),
            ("IPC 420 / BNS 318","Cheating — 7 years"),
            ("IPC 468 / BNS 339","Forgery for Cheating — 7 years"),
        ],
        "Constitutional Articles": [
            ("Article 14","Right to Equality before Law"),
            ("Article 19","Freedom of Speech, Assembly, Movement"),
            ("Article 21","Right to Life and Personal Liberty"),
            ("Article 22","Protection against Arbitrary Arrest"),
            ("Article 32","Right to Constitutional Remedies (Supreme Court)"),
            ("Article 226","High Court Writ Jurisdiction"),
        ],
        "Bail Provisions": [
            ("CrPC 436 / BNSS 478","Bail in Bailable Offences — right, not discretion"),
            ("CrPC 437 / BNSS 480","Non-Bailable — Magistrate's discretion"),
            ("CrPC 438 / BNSS 482","Anticipatory Bail — Sessions Court / High Court"),
            ("CrPC 439 / BNSS 483","Special Powers of High Court and Sessions Court"),
        ],
        "Limitation Periods": [
            ("3 Years","Most civil suits (money, contract, tort)"),
            ("12 Years","Suits on immovable property"),
            ("90 Days","Consumer complaint (extendable)"),
            ("1 Year","Cheque bounce u/s 138 NI Act"),
        ],
    }

    for cat, items in ref_data.items():
        rows = "".join(
            f'<tr><td>{sec}</td><td>{desc}</td></tr>'
            for sec, desc in items
        )
        with st.expander(cat, expanded=False):
            st.markdown(
                f'<table class="ref-table"><tbody>{rows}</tbody></table>',
                unsafe_allow_html=True,
            )

    st.markdown(
        '<p class="info-note">For information only. Consult a qualified advocate for your specific matter.</p>',
        unsafe_allow_html=True,
    )

# ════════════════════════════════════════════════════
# TAB 5 — METRICS
# ════════════════════════════════════════════════════
with tab_metrics:
    st.markdown('<p class="section-eyebrow">Performance</p>', unsafe_allow_html=True)
    st.markdown('<h1 class="section-title">Session Metrics</h1>', unsafe_allow_html=True)

    m = st.session_state.metrics
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    for col, val, lbl in [
        (c1, m["total_queries"],               "Queries"),
        (c2, f"{avg(m['llm_times'])/1000:.1f}s" if m["llm_times"] else "—", "Avg LLM"),
        (c3, f"{avg(m['retrieval_times'])}ms"  if m["retrieval_times"] else "—", "Avg Retrieval"),
        (c4, f"{avg([e['total_ms'] for e in m['query_log']])/1000:.1f}s" if m["query_log"] else "—", "Avg Total"),
        (c5, f"{m['total_tokens_est']:,}",     "~Tokens"),
        (c6, m["total_doc_generations"],       "Docs"),
    ]:
        col.markdown(
            f'<div class="metric-card"><div class="metric-val">{val}</div>'
            f'<div class="metric-lbl">{lbl}</div></div>',
            unsafe_allow_html=True,
        )

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    if not m["query_log"]:
        st.markdown("""
        <div class="empty-state">
          <div class="empty-glyph">📊</div>
          <p class="empty-title">No data yet</p>
          <p class="empty-sub">Use the chatbot or document generator to collect metrics.</p>
        </div>
        """, unsafe_allow_html=True)
    else:
        ca, cb = st.columns(2)

        with ca:
            st.markdown(
                '<div style="background:var(--navy-surface);border:1px solid var(--gold-rule);'
                'border-radius:5px;padding:1.2rem 1.4rem">'
                '<div style="font-family:var(--mono);font-size:0.6rem;font-weight:300;'
                'letter-spacing:0.16em;text-transform:uppercase;color:var(--gold-deep);margin-bottom:10px">'
                'Recent Queries</div>',
                unsafe_allow_html=True,
            )
            for entry in list(reversed(m["query_log"]))[:12]:
                s   = entry["total_ms"] / 1000
                clr = "#5AB87A" if s < 3 else "#C9A84C" if s < 8 else "#BF6B6B"
                ret = f' · 🔍 {entry["retrieval_ms"]}ms' if entry["retrieval_ms"] else ""
                st.markdown(
                    f'<div class="log-row">'
                    f'<span style="color:var(--text-faint)">{entry["time"]}</span>'
                    f'&nbsp;<span style="color:var(--gold)">{entry["model"][:22]}</span>'
                    f'&nbsp;<span style="color:var(--text-faint);font-size:0.63rem">[{entry["type"]}]</span>'
                    f'<span style="color:{clr}">⚡ {s:.1f}s{ret}</span></div>',
                    unsafe_allow_html=True,
                )
            st.markdown('</div>', unsafe_allow_html=True)

        with cb:
            st.markdown(
                '<div style="background:var(--navy-surface);border:1px solid var(--gold-rule);'
                'border-radius:5px;padding:1.2rem 1.4rem">'
                '<div style="font-family:var(--mono);font-size:0.6rem;font-weight:300;'
                'letter-spacing:0.16em;text-transform:uppercase;color:var(--gold-deep);margin-bottom:10px">'
                'Model Usage</div>',
                unsafe_allow_html=True,
            )
            total_u = sum(m["model_usage"].values()) or 1
            for mdl, cnt in sorted(m["model_usage"].items(), key=lambda x: -x[1]):
                pct = round(cnt / total_u * 100)
                st.markdown(
                    f'<div style="margin-bottom:10px">'
                    f'<div style="display:flex;justify-content:space-between;margin-bottom:3px">'
                    f'<span style="font-family:var(--mono);font-size:0.65rem;font-weight:300;color:var(--gold)">{mdl[:32]}</span>'
                    f'<span style="font-family:var(--mono);font-size:0.63rem;font-weight:300;color:var(--text-bright)">{cnt} ({pct}%)</span></div>'
                    f'<div class="perf-bar"><div class="perf-fill" style="width:{pct}%"></div></div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
            st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    if st.button("Reset Metrics"):
        st.session_state.metrics = {
            "total_queries": 0, "total_doc_generations": 0, "total_tokens_est": 0,
            "query_log": [], "llm_times": [], "retrieval_times": [], "model_usage": {},
        }
        st.rerun()
